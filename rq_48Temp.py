# --*-- coding:utf-8 --*--

# Author: Taoyong
# Time: 2019/1/7 12:16

import requests
import json
import pandas as pd
import numpy as np
from pyecharts import Line
from datetime import date, timedelta
from docx import Document
from docx.shared import Inches

ip = "http://10.203.89.55/cimiss-web/api"
userId = "BEGY_XQX_API"
pwd = "xqx_123"
interfaceId = "getSevpWefcRffcByTimeAndStaIDAndValidTime"
elements = "Station_Name,Validtime,TEM"
eleValueRanges = "Prod_Code:SCMOC"
validTimes = "24,27,30,33,36,39,42,45,48"
staIds = "57807,57707,57713,57717,57806,57816,57827"
time = date.today().strftime('%Y%m%d'+'000000')
dataCode = 'SEVP_CHN_WEFC_RFFC_HCB'
dataFormat = 'json'


def getData():
    url=ip+'?userId='+userId+"&pwd="+pwd+"&interfaceId="+interfaceId+"&elements="+elements+"&eleValueRanges="+\
        eleValueRanges+"&validTimes="+validTimes+"&staIds="+staIds+"&time="+time+"&dataCode="+dataCode+"&dataFormat="+dataFormat
    r=requests.get(url).content.decode()
    dic=json.loads(r)
    df=pd.DataFrame(dic['DS'])
    df.set_index('Station_Name',inplace=True)
    df['4℃']=4
    df['7℃']=7
    df['11℃']=11
    return df

# 毕节、六枝曲线图
def plot_pic1(arr1, data1, d):
    line = Line(title='毕节、六枝逐3小时气温预报', title_pos='center', title_top='5%', height=600, width=1000)
    line.add('毕节', arr1, data1.loc['毕节']['TEM'], is_label_show=True, is_smooth=True, line_width=2, mark_point=['max', 'min'],
             mark_point_symbolsize=25, mark_point_symbol='circle', line_color="red")
    line.add('六枝', arr1, data1.loc['六枝']['TEM'], is_label_show=True, is_smooth=True, line_width=2, mark_point=['max', 'min'],
             mark_point_symbolsize=25, mark_point_symbol='circle', line_color="blue")
    line.add('4℃', arr1, data1.loc['毕节']['4℃'], line_type='dashed', line_color="orange")
    line.add('7℃', arr1, data1.loc['毕节']['7℃'], xaxis_name='时间', yaxis_name='温度', is_xaxislabel_align=True,
             is_splitline_show=False, line_type='dashed')
    line.add('11℃', arr1, data1.loc['毕节']['11℃'], xaxis_name='时间', yaxis_name='温度', is_xaxislabel_align=True,
             is_splitline_show=False, line_type='dashed')
    line.render(path=d+'毕节、六枝'+'.gif')
# 安顺、贵阳、都匀曲线图
def plot_pic2(arr2, data2, d):
    line = Line(title='安顺、贵阳、都匀逐3小时气温预报', title_pos='center', title_top='5%', height=600, width=1000)
    line.add('安顺', arr2, data2.loc['安顺']['TEM'], is_label_show=True, is_smooth=True, line_width=2, mark_point=['max', 'min'],
             mark_point_symbolsize=25, mark_point_symbol='circle', line_color="red")
    line.add('贵阳', arr2, data2.loc['贵阳']['TEM'], is_label_show=True, is_smooth=True, line_width=2, mark_point=['max', 'min'],
             mark_point_symbolsize=25, mark_point_symbol='circle', line_color="blue")
    line.add('都匀', arr2, data2.loc['都匀']['TEM'], is_label_show=True, is_smooth=True, line_width=2, mark_point=['max', 'min'],
             mark_point_symbolsize=25, mark_point_symbol='circle', line_color="green")
    line.add('4℃', arr2, data2.loc['安顺']['4℃'], line_type='dashed', line_color="orange")
    line.add('7℃', arr2, data2.loc['安顺']['7℃'], xaxis_name='时间', yaxis_name='温度', is_xaxislabel_align=True,
             is_splitline_show=False, line_type='dashed')
    line.add('11℃', arr2, data2.loc['安顺']['11℃'], xaxis_name='时间', yaxis_name='温度', is_xaxislabel_align=True,
             is_splitline_show=False, line_type='dashed')
    line.render(path=d+'安顺、贵阳、都匀'+'.gif')
# 遵义、播州区曲线图
def plot_pic3(arr3, data3, d):
    line = Line(title='遵义、播州区逐3小时气温预报', title_pos='center', title_top='5%', height=600, width=1000)
    line.add('遵义', arr3, data3.loc['遵义']['TEM'], is_label_show=True, is_smooth=True, line_width=2, mark_point=['max', 'min'],
             mark_point_symbolsize=25, mark_point_symbol='circle', line_color="red")
    line.add('播州区', arr3, data3.loc['播州区']['TEM'], is_label_show=True, is_smooth=True, line_width=2, mark_point=['max', 'min'],
             mark_point_symbolsize=25, mark_point_symbol='circle', line_color="blue")
    line.add('4℃', arr3, data3.loc['遵义']['4℃'], line_type='dashed', line_color="orange")
    line.add('7℃', arr3, data3.loc['遵义']['7℃'], xaxis_name='时间', yaxis_name='温度', is_xaxislabel_align=True,
             is_splitline_show=False, line_type='dashed')
    line.add('11℃', arr3, data3.loc['遵义']['11℃'], xaxis_name='时间', yaxis_name='温度', is_xaxislabel_align=True,
             is_splitline_show=False, line_type='dashed')
    line.render(path=d+'遵义、播州区'+'.gif')
# 打开word模板，插入生成的图片，最后保存为最新的文档
def insert_word(t2):
    document = Document('贵州燃气公司专题服务模板.docx')
    # para = document.paragraphs
    # for p in para:
    #     print(p.text)
    # for p in range(len(document.paragraphs)):
    #     if p>3:
    #         document.paragraphs[p].clear()
    document.add_picture(t2 + '毕节、六枝.gif', width=Inches(7.25))
    document.add_picture(t2 + '安顺、贵阳、都匀.gif', width=Inches(7.25))
    document.add_picture(t2 + '遵义、播州区.gif', width=Inches(7.25))
    document.save('贵州燃气公司专题服务'+t2+'.docx')

def main():

    # 获取昨天、今天、明天日期
    # zt = (date.today() + timedelta(days=-1)).strftime("%Y%m%d")
    jt = date.today().strftime("%Y%m%d")
    tomorrow = (date.today() + timedelta(days=1)).strftime("%d")
    postnatal = (date.today() + timedelta(days=2)).strftime("%d")
    arr = ['08时', '11时', '14时', '17时', '20时', '23时', '02时', '05时', '08时']
    # 处理日期隔天
    for i in range(0, len(arr)):
        if i < 6:
            arr[i] = tomorrow + '日' + arr[i]
        else:
            arr[i] = postnatal + '日' + arr[i]
    # 读取数据
    data = getData()
    # data = pd.read_table('data_3h.txt', header=None, skiprows=1, encoding='utf-8', index_col=0)
    print(data)

    plot_pic1(arr, data, jt)
    plot_pic2(arr, data, jt)
    plot_pic3(arr, data, jt)
    insert_word(jt)

if __name__ == "__main__":
    main()
